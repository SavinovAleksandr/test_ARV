# -*- coding: utf-8 -*-
"""
Формирование отчёта «Приложение. Результат в графическом виде.docx» с вставкой рисунков.
"""
import os
from typing import List, Tuple

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


def write_result_docx(figures: List[Tuple[str, str]], out_path: str) -> None:
    """
    Создать docx: каждый элемент figures — (подпись, путь к PNG).
    """
    if Document is None:
        raise RuntimeError("Установите python-docx: pip install python-docx")
    doc = Document()
    for caption, png_path in figures:
        if not os.path.isfile(png_path):
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(png_path, width=Inches(5.5))
        p2 = doc.add_paragraph(caption)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(out_path)
